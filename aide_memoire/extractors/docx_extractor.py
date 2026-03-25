"""Word document (.docx) content extractor."""

from pathlib import Path

from docx import Document

from aide_memoire.extractors.base import BaseExtractor
from aide_memoire.models import ContentBlock, ContentType, ExtractedDocument


class DocxExtractor(BaseExtractor):
    def extract(self, file_path: Path, output_dir: Path | None = None) -> ExtractedDocument:
        doc_word = Document(str(file_path))
        doc = ExtractedDocument(source_path=file_path, title=file_path.stem)

        # Extract paragraphs
        for i, para in enumerate(doc_word.paragraphs, 1):
            text = para.text.strip()
            if text:
                doc.blocks.append(ContentBlock(
                    content_type=ContentType.TEXT,
                    text=text,
                    source_file=file_path.name,
                    source_location=f"Paragraph {i}",
                ))

        # Extract tables
        for i, table in enumerate(doc_word.tables, 1):
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            table_text = "\n".join(rows)
            if table_text.strip():
                doc.blocks.append(ContentBlock(
                    content_type=ContentType.TABLE,
                    text=table_text,
                    source_file=file_path.name,
                    source_location=f"Table {i}",
                ))

        return doc
